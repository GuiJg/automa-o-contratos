from docx import Document
from docx.shared import Inches
from datetime import datetime
import pandas as pd
import locale
import os

# Configure a localidade para português do Brasil
locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')  # ou 'pt_BR' dependendo do seu sistema

# Leitura do arquivo Excel
tabela = pd.read_excel("informações.xlsx")

for linha in tabela.index:
    # Criar um novo documento
    documento = Document("contrato.docx")

    # Obter dados da linha atual
    nome = tabela.loc[linha, "Nome"]
    rg = tabela.loc[linha, "Rg"]
    cpf = tabela.loc[linha, "Cpf"] 
    orgao = tabela.loc[linha, "Orgao"]
    nomemae = tabela.loc[linha, "NomeMae"]
    nomepai = tabela.loc[linha, "NomePai"]
    endereco = tabela.loc[linha, "Endereco"]
    matricula = tabela.loc[linha, "Matricula"]
    cidade = tabela.loc[linha, "Cidade"]
    email = tabela.loc[linha, "Email"]
    telefone = tabela.loc[linha, "Telefone"]
    nomedepen = tabela.loc[linha, "NomeDepen"]
    nascimento = tabela.loc[linha, "Nascimento"]
    parentesco = tabela.loc[linha, "Parentesco"]

    # Dicionário de referências
    referencias = {
        "XXXX": nome,
        "RRRR": rg,
        "FFFF": cpf,
        "OOOO": orgao,
        "XXXM": nomemae,
        "XXXP": nomepai,
        "EEEE": endereco,
        "TTTT": matricula,
        "CCCC": cidade,
        "YYYY": email,
        "ZZZZ": telefone,
        "XXXD": nomedepen,
        "NNNN": nascimento,
        "PPPP": parentesco,
        "DD": str(datetime.now().day),
        "MM": datetime.now().strftime("%B"),
        "AAAA": str(datetime.now().year),
    }

    # Substituir as referências no documento
    for paragrafo in documento.paragraphs:
        for codigo in referencias:
            valor = referencias[codigo]
            paragrafo.text = paragrafo.text.replace(codigo, str(valor))

    # Adicionar imagem ao documento
    image_path = "logo.jpg"  # Substitua pelo caminho real da sua imagem
    if os.path.exists(image_path):
        documento.add_picture(image_path, width=Inches(2))

    # Salvar o documento com um nome específico
    documento.save(f"Filiação - {nome} - {cidade}.docx")

# Restaurar a localidade para a original (opcional)
locale.setlocale(locale.LC_TIME, '')
